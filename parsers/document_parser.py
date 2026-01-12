"""
Document parser for extracting text from PDF and DOCX files.
"""

import io
from typing import BinaryIO

import PyPDF2
from docx import Document


def extract_text_from_pdf(file: BinaryIO) -> str:
    """
    Extract text content from a PDF file.
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text_parts = []

        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file: BinaryIO) -> str:
    """
    Extract text content from a DOCX file.
    """
    try:
        doc = Document(file)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def parse_document(file, filename: str) -> str:
    """
    Auto-detect file format and extract text.
    """
    filename_lower = filename.lower()

    if hasattr(file, "seek"):
        file.seek(0)

    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx")
    elif filename_lower.endswith(".txt"):
        content = file.read()
        if isinstance(content, bytes):
            return content.decode("utf-8", errors="ignore")
        return content
    else:
        raise ValueError(
            f"Unsupported file format: {filename}. Supported formats: PDF, DOCX, TXT"
        )
